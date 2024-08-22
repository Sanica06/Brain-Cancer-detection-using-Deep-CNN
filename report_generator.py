import os
os.environ["SM_FRAMEWORK"] = "tf.keras"
import numpy as np
import tensorflow as tf
import segmentation_models as sm
from PIL import Image, ImageDraw
from docx import Document
from docx.shared import Inches
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
import cv2

BACKBONE = 'resnext50'
preprocess_input = sm.get_preprocessing(BACKBONE)

def preprocess_image(image_path):
    """ Preprocess the image for prediction. """
    img = cv2.imread(image_path)
    img = cv2.resize(img, (256, 256))  
    img = preprocess_input(img) 
    img_array = np.expand_dims(img, axis=0) 
    return Image.fromarray(cv2.cvtColor(img, cv2.COLOR_BGR2RGB)), img_array

def predict_tumor(model, image_array, threshold=0.5):
    """ Predict whether there is a tumor. """
    pr_mask = model.predict(image_array)    
    if isinstance(pr_mask, list):
        pr_mask = pr_mask[-1]  
    pr_mask_binary = (pr_mask > threshold).astype(np.uint8)
    
    return pr_mask_binary

def highlight_tumor(original_img, segmentation_map):
    """ Highlight the tumor region on the image with a bounding box. """
    segmentation_map_resized = cv2.resize(segmentation_map[0], (original_img.size[0], original_img.size[1]), interpolation=cv2.INTER_NEAREST)

    if segmentation_map_resized.ndim != 2:
        raise ValueError("The segmentation map should be a 2D array.")

    y_indices, x_indices = np.where(segmentation_map_resized == 1)
    if len(x_indices) == 0 or len(y_indices) == 0:
        return original_img

    x_min, x_max = x_indices.min(), x_indices.max()
    y_min, y_max = y_indices.min(), y_indices.max()

    overlay = original_img.copy()
    draw = ImageDraw.Draw(overlay)
    draw.rectangle([x_min, y_min, x_max, y_max], outline="red", width=3)
    return overlay

def create_report(image_name, has_tumor, highlighted_img, report_path, original_img):
    """ Create a report in a Word document. """
    doc = Document()
    doc.add_heading('Tumor Detection Report', 0)
    doc.add_paragraph(f'Image: {image_name}')
    doc.add_paragraph(f'Tumor detected: {"Yes" if has_tumor else "No"}')
    
    original_img_path = '/tmp/original_image.png'
    original_img.save(original_img_path)
    doc.add_paragraph('Original Image:')
    doc.add_picture(original_img_path, width=Inches(5.0))

    if has_tumor:
        highlighted_img_path = '/tmp/highlighted_image.png'
        highlighted_img.save(highlighted_img_path)
        doc.add_paragraph('Highlighted Image:')
        doc.add_picture(highlighted_img_path, width=Inches(5.0))
    
    doc.save(report_path)

def create_pdf_report(image_name, has_tumor, highlighted_img, report_path, original_img):
    """ Create a report in a PDF document. """
    c = canvas.Canvas(report_path, pagesize=A4)
    width, height = A4

    c.setFont("Helvetica-Bold", 20)
    c.drawString(100, height - 100, "Tumor Detection Report")

    c.setFont("Helvetica", 12)
    c.drawString(100, height - 140, f"Image: {image_name}")
    c.drawString(100, height - 160, f"Tumor detected: {'Yes' if has_tumor else 'No'}")
    
    original_img_path = '/tmp/original_image.png'
    original_img.save(original_img_path)
    c.drawString(100, height - 180, "Original Image:")
    c.drawImage(original_img_path, 100, height - 400, width=400, height=200)

    if has_tumor:
        highlighted_img_path = '/tmp/highlighted_image.png'
        highlighted_img.save(highlighted_img_path)
        c.drawString(100, height - 420, "Highlighted Image:")
        c.drawImage(highlighted_img_path, 100, height - 700, width=400, height=200)
    
    c.save()

def process_and_generate_reports(image_path, model, docx_report_path='tumor_report.docx', pdf_report_path='tumor_report.pdf'):
    """ Process the image, predict tumor, and generate both Word and PDF reports. """
    original_img, image_array = preprocess_image(image_path)
    has_tumor = predict_tumor(model, image_array)

    if has_tumor.any():
        highlighted_img = highlight_tumor(original_img, has_tumor)
    else:
        highlighted_img = None

    create_report(image_path, has_tumor.any(), highlighted_img, report_path=docx_report_path, original_img=original_img)
    create_pdf_report(image_path, has_tumor.any(), highlighted_img, report_path=pdf_report_path, original_img=original_img)
